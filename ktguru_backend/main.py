"""
KT Guru — FastAPI RAG API for knowledge transfer.
"""
from __future__ import annotations

import io
import logging
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document as DocxDocument

from fastapi import Depends, FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from config import settings
from database import get_db
from models import Document, DocumentChunk, Issue, UnansweredQuestion
from schemas import AskRequest, AskResponse
from services.ai_service import AIService
import os
import certifi

os.environ["SSL_CERT_FILE"] = certifi.where()
os.environ["REQUESTS_CA_BUNDLE"] = certifi.where()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

ai_service: AIService | None = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    global ai_service
    ai_service = AIService()
    ai_service.initialize()
    logger.info("KT Guru API startup complete (FAISS dir: %s)", settings.FAISS_DIR)
    yield
    ai_service = None


app = FastAPI(
    title="KT Guru",
    description="Knowledge Transfer RAG — document upload and Q&A",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origin_list,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
def root() -> dict[str, str]:
    return {"service": "KT Guru API", "docs": "/docs"}


MAX_UPLOAD_BYTES = 5 * 1024 * 1024

_UTF8_TEXT_EXTS = frozenset(
    {"txt", "py", "java", "js", "xml", "yaml", "json", "html"}
)
_CSV_EXT = "csv"
_DOCX_EXT = "docx"
_TABULAR_EXTS = frozenset({_CSV_EXT, "xlsx", "xls"})
_ALLOWED_UPLOAD_EXTS = _UTF8_TEXT_EXTS | _TABULAR_EXTS | {_DOCX_EXT}

# Structured KT columns (CSV / Excel)
_ISSUE_COLUMNS: tuple[str, ...] = (
    "Subject",
    "Raised By",
    "Raised On",
    "Resolution Time",
    "Status",
    "Resolution",
)


def _file_ext(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def _scalar_str(val: Any) -> str:
    if pd.isna(val):
        return ""
    return str(val).strip()


def _read_tabular_dataframe(filename: str, raw: bytes) -> pd.DataFrame:
    ext = _file_ext(filename)
    buf = io.BytesIO(raw)
    if ext == _CSV_EXT:
        return pd.read_csv(buf)
    if ext == "xlsx":
        return pd.read_excel(buf, engine="openpyxl")
    if ext == "xls":
        return pd.read_excel(buf, engine="xlrd")
    raise ValueError("Unsupported tabular file type")


def _normalize_and_validate_issues_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    missing = [c for c in _ISSUE_COLUMNS if c not in df.columns]
    if missing:
        raise ValueError("Missing required columns: " + ", ".join(missing))
    df = df[list(_ISSUE_COLUMNS)].copy()
    if len(df) == 0:
        raise ValueError("No data rows found in file.")
    df["Raised On"] = pd.to_datetime(df["Raised On"], errors="coerce")
    if df["Raised On"].isna().any():
        n_bad = int(df["Raised On"].isna().sum())
        raise ValueError(
            f"'Raised On' must contain valid dates for all rows ({n_bad} invalid or empty)."
        )
    return df


def _issue_row_rich_context(row: pd.Series) -> str:
    subj = _scalar_str(row["Subject"]) or "(no subject)"
    raised_by = _scalar_str(row["Raised By"])
    status = _scalar_str(row["Status"])
    res_time = _scalar_str(row["Resolution Time"])
    resolution = _scalar_str(row["Resolution"])
    ro = row["Raised On"]
    if pd.isna(ro):
        date_disp = ""
    else:
        ts = pd.Timestamp(ro)
        date_disp = ts.strftime("%Y-%m-%d %H:%M:%S")
    return (
        f"Subject: {subj} | Status: {status} | Raised By: {raised_by} | "
        f"Date: {date_disp} | Resolution Time: {res_time} | Resolution: {resolution}"
    )


def _issue_row_to_model(document_id: int, row: pd.Series) -> Issue:
    ro = row["Raised On"]
    dt = pd.Timestamp(ro).to_pydatetime() if not pd.isna(ro) else None
    subj = _scalar_str(row["Subject"]) or "(no subject)"
    rb = _scalar_str(row["Raised By"]) or "—"
    rt = _scalar_str(row["Resolution Time"])
    st = _scalar_str(row["Status"])
    res = _scalar_str(row["Resolution"])
    return Issue(
        document_id=document_id,
        subject=subj,
        raised_by=rb,
        raised_on=dt,
        resolution_time=rt or None,
        status=st or None,
        resolution=res or None,
    )


def _extract_text_from_file(filename: str, raw: bytes) -> str:
    """
    Decode or parse uploaded bytes to plain text for the same
    RecursiveCharacterTextSplitter + FAISS path as the original .txt flow.
    """
    ext = _file_ext(filename)

    if ext in _UTF8_TEXT_EXTS:
        try:
            return raw.decode("utf-8", errors="strict")
        except UnicodeDecodeError as e:
            raise HTTPException(
                status_code=400,
                detail="File must be valid UTF-8 text",
            ) from e

    if ext == _DOCX_EXT:
        buf = io.BytesIO(raw)
        try:
            document = DocxDocument(buf)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail="Could not parse Word document",
            ) from e
        parts: list[str] = []
        for p in document.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        return "\n\n".join(parts) if parts else ""

    raise HTTPException(status_code=400, detail="Unsupported file type")


def get_ai() -> AIService:
    if ai_service is None:
        raise HTTPException(status_code=503, detail="AI service not ready")
    return ai_service


async def get_ai_dep(
    _req: Request,
) -> AIService:
    s = get_ai()
    return s


def _split_document(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        is_separator_regex=False,
    )
    return splitter.split_text(text)


@app.post("/api/upload")
async def upload_document(
    file: UploadFile = File(..., description="Document file to index (text, code, CSV, Excel, or DOCX)"),
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_dep),
) -> dict[str, Any]:
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    ext = _file_ext(file.filename)
    if ext not in _ALLOWED_UPLOAD_EXTS:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    raw = await file.read()
    if len(raw) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5 MB)")

    doc = Document(filename=file.filename)
    db.add(doc)
    await db.flush()
    await db.refresh(doc)

    meta_batch: list[dict[str, Any]] = []
    text_batch: list[str] = []
    chunk_count: int

    if ext in _TABULAR_EXTS:
        try:
            df = _read_tabular_dataframe(file.filename, raw)
            df = _normalize_and_validate_issues_df(df)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e)) from e
        except Exception:
            logger.exception("Tabular ingest failed")
            raise HTTPException(
                status_code=400,
                detail="The file is corrupted or could not be read. Required columns: "
                + ", ".join(_ISSUE_COLUMNS),
            ) from None

        chunk_index = 0
        for _, row in df.iterrows():
            db.add(_issue_row_to_model(doc.id, row))
            db.add(
                DocumentChunk(
                    document_id=doc.id,
                    chunk_text=_issue_row_rich_context(row),
                    chunk_index=chunk_index,
                )
            )
            chunk_index += 1
        chunk_count = chunk_index
        await db.flush()
    else:
        try:
            text = _extract_text_from_file(file.filename, raw)
        except HTTPException:
            raise
        except Exception as e:
            logger.exception("Document extraction failed")
            raise HTTPException(
                status_code=400,
                detail="The file is corrupted or could not be read.",
            ) from e

        chunks = _split_document(text)
        if not chunks:
            raise HTTPException(status_code=400, detail="No text content to index after splitting")

        for i, chunk_text in enumerate(chunks):
            db.add(
                DocumentChunk(
                    document_id=doc.id,
                    chunk_text=chunk_text,
                    chunk_index=i,
                )
            )
        chunk_count = len(chunks)
        await db.flush()

    res = await db.execute(
        select(DocumentChunk)
        .where(DocumentChunk.document_id == doc.id)
        .order_by(DocumentChunk.chunk_index)
    )
    row_chunks = list(res.scalars().all())
    for ch in row_chunks:
        text_batch.append(ch.chunk_text)
        meta_batch.append(
            {
                "filename": file.filename,
                "document_id": str(doc.id),
                "chunk_id": str(ch.id),
                "chunk_index": str(ch.chunk_index or 0),
            }
        )

    await db.commit()

    await ai.add_embedded_chunks(text_batch, meta_batch)

    return {
        "status": "ok",
        "document_id": doc.id,
        "filename": file.filename,
        "chunks": chunk_count,
    }


@app.post("/api/ask", response_model=AskResponse)
async def ask(
    body: AskRequest,
    db: AsyncSession = Depends(get_db),
    ai: AIService = Depends(get_ai_dep),
) -> AskResponse:
    if not settings.GROQ_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="GROQ_API_KEY is not configured on the server",
        )

    q = body.query.strip()
    if not q:
        raise HTTPException(status_code=400, detail="query must not be empty")

    answer, sources, log_unanswered = await ai.ask(q)

    if log_unanswered:
        uq = UnansweredQuestion(
            question=q,
            status="UNRESOLVED",
        )
        db.add(uq)
        await db.commit()

    return AskResponse(answer=answer, sources=sources)


@app.get("/api/health")
async def health(
    db: AsyncSession = Depends(get_db),
) -> dict[str, str]:
    await db.execute(select(Document).limit(1))
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
    )
